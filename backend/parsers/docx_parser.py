"""
Word (.docx) 文档解析器
提取段落、表格、标题等内容
"""
from docx import Document as DocxDocument
from docx.table import Table
from typing import List
import os


def parse_docx(file_path: str) -> dict:
    """
    解析 Word 文档，提取文本和表格内容
    
    参数:
        file_path: .docx 文件路径
    
    返回:
        {
            "title": 文档标题,
            "sections": [{"heading": 标题, "level": 级别, "content": 内容}],
            "tables": [[行数据]],
            "raw_content": 纯文本内容
        }
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    doc = DocxDocument(file_path)
    title = ""
    sections = []
    current_section = None
    current_content = []
    tables = []
    all_text = []

    # 提取段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text.append(text)

        # 判断是否为标题
        if para.style and para.style.name.startswith("Heading"):
            # 保存上一个 section
            if current_section:
                current_section["content"] = "\n".join(current_content).strip()
                sections.append(current_section)

            level = 1
            try:
                level = int(para.style.name.replace("Heading", "").strip())
            except ValueError:
                level = 1

            if not title and level == 1:
                title = text

            current_section = {
                "heading": text,
                "level": level,
                "content": ""
            }
            current_content = []
        else:
            current_content.append(text)

    # 保存最后一个 section
    if current_section:
        current_section["content"] = "\n".join(current_content).strip()
        sections.append(current_section)

    # 提取表格
    for table in doc.tables:
        table_data = _extract_table(table)
        if table_data:
            tables.append(table_data)

    if not title:
        title = os.path.basename(file_path)

    return {
        "title": title,
        "sections": sections,
        "tables": tables,
        "raw_content": "\n".join(all_text)
    }


def _extract_table(table: Table) -> List[List[str]]:
    """提取表格数据"""
    result = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        if any(row_data):  # 跳过全空行
            result.append(row_data)
    return result
